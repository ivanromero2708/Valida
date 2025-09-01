# render_todos_los_parametros.py
# -*- coding: utf-8 -*-
# Requisitos:
#   pip install docxtpl python-docx

from __future__ import annotations

import re
import time
from dataclasses import dataclass, asdict
from datetime import datetime
from math import sqrt, isfinite
from pathlib import Path
from statistics import mean, stdev
from typing import Any, Dict, List, Optional, Tuple

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# Configuración general
# =========================
BASE = Path(
    r"D:\Users\aespinosa\OneDrive - Grupo Procaps\NUEVAS PLATAFORMAS\D&R 4.0\PROYECTOS_2025_PROCAPS\DESARROLLOS\VALIDACIONES\pruebas"
)
TPL = BASE / "templates"
OUT = BASE / "output"
OUT.mkdir(parents=True, exist_ok=True)


# =========================
# Utilidades genéricas
# =========================
def fmt2(x: Optional[float]) -> str:
    return "" if x is None else f"{float(x):.2f}"


def safe_mean(vals: List[Optional[float]]) -> Optional[float]:
    v = [float(x) for x in vals if x is not None]
    return None if not v else float(mean(v))


def robust_save(
    docx: DocxTemplate, out_path: Path, reintentos: int = 3, espera: float = 0.7
) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    for i in range(reintentos):
        try:
            docx.save(str(out_path))
            return out_path
        except PermissionError:
            if i == reintentos - 1:
                alt = out_path.with_name(
                    out_path.stem + f"_{datetime.now():%Y%m%d_%H%M%S}_tmp.docx"
                )
                docx.save(str(alt))
                return alt
            time.sleep(espera)


def render_docxtpl(
    template_path: Path, contexto: Dict[str, Any], out_path: Path
) -> Optional[Path]:
    if not template_path.exists():
        print(f"⚠ Plantilla no encontrada, se omite:\n  {template_path}")
        return None
    if template_path.suffix.lower() != ".docx":
        print(f"⚠ La plantilla no es .docx, se omite:\n  {template_path}")
        return None
    tpl = DocxTemplate(str(template_path))
    try:
        missing = tpl.get_undeclared_template_variables(contexto)
        if missing:
            print("  · Aviso: variables no declaradas en contexto:", sorted(missing))
    except Exception:
        pass
    tpl.render(contexto)
    return robust_save(tpl, out_path)


# ============================================================
# SECCIÓN 1 — LINEALIDAD (multi-activos, SIN vMerge post-render)
# ============================================================
@dataclass
class Replica:
    concentracion: float  # mg/mL
    area_pico: float
    factor_respuesta: float  # típicamente area/concentración


def rsd_percent(vals: List[float]) -> float:
    if len(vals) < 2:
        return 0.0
    mu = mean(vals)
    if mu == 0:
        return 0.0
    return 100.0 * (stdev(vals) / mu)


def regresion_simple(xs: List[float], ys: List[float]) -> Dict[str, float]:
    """
    y = m*x + b (mínimos cuadrados), correlación r y r^2.
    Recomendado: usar y promedio por nivel (área promedio) vs x (concentración).
    """
    n = len(xs)
    if n < 2:
        return {"slope": 0.0, "intercept": 0.0, "r": 0.0, "r2": 0.0}

    sx = sum(xs)
    sy = sum(ys)
    sxy = sum(x * y for x, y in zip(xs, ys))
    sx2 = sum(x * x for x in xs)
    sy2 = sum(y * y for y in ys)

    denom = n * sx2 - sx * sx
    if denom == 0:
        m = 0.0
        b = sum(ys) / n
    else:
        m = (n * sxy - sx * sy) / denom
        b = (sy - m * sx) / n

    denom_r = sqrt((n * sx2 - sx * sx) * (n * sy2 - sy * sy))
    if denom_r == 0:
        r = 0.0
    else:
        r = (n * sxy - sx * sy) / denom_r
    r2 = r * r
    if not isfinite(r):
        r, r2 = 0.0, 0.0
    return {"slope": m, "intercept": b, "r": r, "r2": r2}


def construir_contexto_linealidad_multi_activos(
    activos_in: List[Dict[str, Any]],
    criterios_por_defecto: Dict[str, Any] | None = None,
) -> Dict[str, Any]:
    """
    Estructura esperada por activo:
    {
      "nombre": "API X",
      "linealidad_sistema": [
        {"nivel": 1, "replicas": [Replica|dict, ...]},
        ...
      ],
      # opcional:
      "criterios": {
         "r_min": 0.998,
         "rsd_max": 2.0,
         "pct_intercepto_max": 2.0,   # opcional
         "x_ref_100pct": 0.3000       # opcional (mg/mL) para %intercepto
      }
    }
    """
    if criterios_por_defecto is None:
        criterios_por_defecto = {"r_min": 0.998, "rsd_max": 2.0}

    activos_out = []

    for a in activos_in:
        nombre = str(a["nombre"])

        # Normaliza réplicas (dataclass -> dict)
        lineas = []
        for d in a["linealidad_sistema"]:
            lineas.append(
                {
                    "nivel": d["nivel"],
                    "replicas": [
                        asdict(r) if not isinstance(r, dict) else r
                        for r in d["replicas"]
                    ],
                }
            )

        # Criterios (defaults + override por activo)
        crit = {**criterios_por_defecto, **a.get("criterios", {})}

        # RSD% global de factores de respuesta (todas las réplicas)
        factores_global = [
            float(rep["factor_respuesta"])
            for linea in lineas
            for rep in linea["replicas"]
        ]
        rsd_factor = rsd_percent(factores_global)

        # Regresión: área promedio por nivel vs concentración
        xs = [float(linea["replicas"][0]["concentracion"]) for linea in lineas]
        y_means = [
            float(mean([float(rep["area_pico"]) for rep in linea["replicas"]]))
            for linea in lineas
        ]
        reg = regresion_simple(xs, y_means)
        m, b, r, r2 = reg["slope"], reg["intercept"], reg["r"], reg["r2"]

        # %Intercepto relativo a ŷ(100%) (opcional)
        x_ref = crit.get("x_ref_100pct", max(xs) if xs else 1.0)
        y_hat_ref = m * x_ref + b
        pct_intercepto = (abs(b) / y_hat_ref * 100.0) if y_hat_ref != 0 else 0.0

        # Comparaciones (Cumple/No cumple)
        cumple_r = r >= float(crit["r_min"])
        cumple_rsd = rsd_factor <= float(crit["rsd_max"])
        cumple_pct = None
        if "pct_intercepto_max" in crit:
            cumple_pct = pct_intercepto <= float(crit["pct_intercepto_max"])

        cumple_global = (
            (cumple_r and cumple_rsd)
            if (cumple_pct is None)
            else (cumple_r and cumple_rsd and bool(cumple_pct))
        )

        activos_out.append(
            {
                "nombre": nombre,
                "linealidad_sistema": lineas,  # el template iterará con {%tr%}
                "criterios": crit,
                "metrics": {
                    "rsd_factor": rsd_factor,
                    "pendiente": m,
                    "intercepto": b,
                    "r": r,
                    "r2": r2,
                    "porcentaje_intercepto": pct_intercepto,
                    "cumple_global": cumple_global,
                },
            }
        )

    return {"activos": activos_out}


def render_linealidad():
    template = TPL / "linealidad_multi_activos_tr.docx"  # <- PLANTILLA CORRECTA
    out = OUT / "linealidad_multi_renderizada.docx"

    # Demo (ajusta a tus datos reales)
    activos_demo = [
        {
            "nombre": "TuActivo A",
            "linealidad_sistema": [
                {
                    "nivel": 1,
                    "replicas": [
                        asdict(Replica(0.1000, 1230.5, 12.35)),
                        asdict(Replica(0.1000, 1228.1, 12.32)),
                        asdict(Replica(0.1000, 1236.9, 12.38)),
                    ],
                },
                {
                    "nivel": 2,
                    "replicas": [
                        asdict(Replica(0.2000, 2451.2, 12.26)),
                        asdict(Replica(0.2000, 2460.3, 12.30)),
                        asdict(Replica(0.2000, 2448.7, 12.24)),
                    ],
                },
                {
                    "nivel": 3,
                    "replicas": [
                        asdict(Replica(0.3000, 3665.4, 12.21)),
                        asdict(Replica(0.3000, 3659.8, 12.20)),
                        asdict(Replica(0.3000, 3671.2, 12.24)),
                    ],
                },
                {
                    "nivel": 4,
                    "replicas": [
                        asdict(Replica(0.4000, 4890.1, 12.23)),
                        asdict(Replica(0.4000, 4902.6, 12.26)),
                        asdict(Replica(0.4000, 4881.7, 12.20)),
                    ],
                },
                {
                    "nivel": 5,
                    "replicas": [
                        asdict(Replica(0.5000, 6123.0, 12.25)),
                        asdict(Replica(0.5000, 6132.5, 12.27)),
                        asdict(Replica(0.5000, 6111.9, 12.22)),
                    ],
                },
            ],
            "criterios": {
                "r_min": 0.998,
                "rsd_max": 2.0,
                "pct_intercepto_max": 2.0,
                "x_ref_100pct": 0.3000,
            },
        },
        {
            "nombre": "TuActivo B",
            "linealidad_sistema": [
                {
                    "nivel": 1,
                    "replicas": [
                        asdict(Replica(0.1000, 1230.5, 12.35)),
                        asdict(Replica(0.1000, 1228.1, 12.32)),
                        asdict(Replica(0.1000, 1236.9, 12.38)),
                    ],
                },
                {
                    "nivel": 2,
                    "replicas": [
                        asdict(Replica(0.2000, 2451.2, 12.26)),
                        asdict(Replica(0.2000, 2460.3, 12.30)),
                        asdict(Replica(0.2000, 2448.7, 12.24)),
                    ],
                },
                {
                    "nivel": 3,
                    "replicas": [
                        asdict(Replica(0.3000, 3665.4, 12.21)),
                        asdict(Replica(0.3000, 3659.8, 12.20)),
                        asdict(Replica(0.3000, 3671.2, 12.24)),
                    ],
                },
                {
                    "nivel": 4,
                    "replicas": [
                        asdict(Replica(0.4000, 4890.1, 12.23)),
                        asdict(Replica(0.4000, 4902.6, 12.26)),
                        asdict(Replica(0.4000, 4881.7, 12.20)),
                    ],
                },
                {
                    "nivel": 5,
                    "replicas": [
                        asdict(Replica(0.5000, 6123.0, 12.25)),
                        asdict(Replica(0.5000, 6132.5, 12.27)),
                        asdict(Replica(0.5000, 6111.9, 12.22)),
                    ],
                },
            ],
            # sin "criterios" -> usa defaults
        },
    ]

    ctx = construir_contexto_linealidad_multi_activos(activos_demo)
    p = render_docxtpl(template, ctx, out)
    if p:
        print("✔ Linealidad (multi):", p)


# ============================================================
# SECCIÓN 2 — EXACTITUD (multi-activo)
# ============================================================
def _bloque_exactitud_por_niveles(
    niveles: List[Dict[str, Any]], rango_aceptacion: Tuple[float, float]
) -> List[Dict[str, Any]]:
    lo, hi = rango_aceptacion
    criterio_txt = f"El porcentaje de recuperación promedio por cada nivel debe estar entre {lo:.1f}% y {hi:.1f}%."
    salida: List[Dict[str, Any]] = []
    for item in niveles:
        nivel = str(item["nivel"])
        reps = [float(x) for x in item["replicas"]]
        if not reps:
            raise ValueError(f"El nivel '{nivel}' no tiene réplicas.")
        prom = float(mean(reps))
        cumple = lo <= prom <= hi
        salida.append(
            {
                "nivel": nivel,
                "replicas": [{"recuperacion": v} for v in reps],
                "recuperacion_promedio": prom,
                "criterio": criterio_txt,
                "conclusion": "Cumple" if cumple else "No cumple",
                "cumple_nivel": cumple,
            }
        )
    return salida


def construir_contexto_exactitud_multi(
    activos: List[Dict[str, Any]],
    rango_aceptacion_por_defecto: Tuple[float, float] = (98.0, 102.0),
) -> Dict[str, Any]:
    lista = []
    for a in activos:
        nombre = str(a["nombre"])
        niveles = a["niveles"]
        rango = tuple(a.get("rango_aceptacion", rango_aceptacion_por_defecto))  # type: ignore
        exacts = _bloque_exactitud_por_niveles(niveles, rango)
        fuera = [e for e in exacts if not e["cumple_nivel"]]
        conclusion_global = "Cumple" if not fuera else "No cumple"
        lista.append(
            {
                "nombre": nombre,
                "exactitude_del_metodo": exacts,
                "conclusion_global": conclusion_global,
                "niveles_fuera": [e["nivel"] for e in fuera],
                "rango_aceptacion": {"lo": rango[0], "hi": rango[1]},
            }
        )
    return {
        "activos": lista,
        "activos_titulo": ", ".join([a["nombre"] for a in lista]) if lista else "",
        "fecha_render": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def render_exactitud_multi():
    template = TPL / "exactitud_multi.docx"
    out = OUT / "exactitud_multi_renderizada.docx"
    activos_demo = [
        {
            "nombre": "Compuesto A",
            "niveles": [
                {"nivel": "80%", "replicas": [98.3, 99.1, 97.9]},
                {"nivel": "100%", "replicas": [100.2, 99.8, 100.4]},
                {"nivel": "120%", "replicas": [101.4, 101.9, 100.7]},
            ],
            "rango_aceptacion": (98.0, 102.0),
        },
        {
            "nombre": "Compuesto B",
            "niveles": [
                {"nivel": "80%", "replicas": [96.8, 97.1, 97.4]},
                {"nivel": "100%", "replicas": [98.9, 99.0, 99.2]},
                {"nivel": "120%", "replicas": [102.5, 103.1, 102.9]},
            ],
        },
    ]
    ctx = construir_contexto_exactitud_multi(activos_demo)
    p = render_docxtpl(template, ctx, out)
    if p:
        print("✔ Exactitud multi:", p)


# ============================================================
# SECCIÓN 3 — MATERIALES / INSUMOS
# ============================================================
def context_materiales_demo() -> Dict[str, Any]:
    return {
        "muestra_utilizadas": [
            {
                "nombre": "Caféina 10 mg",
                "codigo": "CAF-10",
                "lote": "L2401-01",
                "codigo_interno_cim": "CIM-0001",
            },
            {
                "nombre": "Ácido ascórbico 50 mg",
                "codigo": "AA-50",
                "lote": "A2503-02",
                "codigo_interno_cim": "CIM-0015",
            },
        ],
        "estandar_utilizados": [
            {
                "nombre": "USP Caffeine RS",
                "fabricante": "USP",
                "lote": "R1234",
                "numero_parte": "123-CAF",
                "codigo_identificacion": "USP-CAF-RS",
                "concentracion": "1000 µg/mL",
                "vencimiento": "26-12-31",
            },
            {
                "nombre": "Ascorbic Acid CRS",
                "fabricante": "Ph.Eur.",
                "lote": "C5678",
                "numero_parte": "AA-CRS-5678",
                "codigo_identificacion": "EP-AA-CRS",
                "concentracion": "500 µg/mL",
                "vencimiento": "27-06-30",
            },
        ],
        "reactivo_utilizados": [
            {
                "nombre": "Metanol LC-MS",
                "fabricante": "Merck",
                "lote": "M2024-07",
                "numero_parte": "1.06007.2500",
                "vencimiento": "27-07-01",
            },
            {
                "nombre": "Acetonitrilo HPLC",
                "fabricante": "J.T.Baker",
                "lote": "ACN-2025-01",
                "numero_parte": "9829-03",
                "vencimiento": "28-01-15",
            },
        ],
        "materiales_utilizados": [
            {
                "nombre": "Filtro PTFE 0.22 µm",
                "fabricante": "Millipore",
                "numero_parte": "SLGV033RS",
                "lote": "F2025-02",
            },
            {
                "nombre": "Vial ámbar 2 mL",
                "fabricante": "Agilent",
                "numero_parte": "5182-0715",
                "lote": "V2024-12",
            },
        ],
        "equipos_utilizados": [
            {
                "nombre": "HPLC 1260 Infinity",
                "consecutivo": "EQ-00123",
                "fabricante": "Agilent",
                "modelo": "G1311B",
                "serial": "US12345678",
                "prox_actividad": "26-11",
            },
            {
                "nombre": "Balanza analítica",
                "consecutivo": "EQ-00456",
                "fabricante": "Mettler Toledo",
                "modelo": "XS205",
                "serial": "MT987654321",
                "prox_actividad": "26-09",
            },
        ],
        "columna_utilizada": [
            {
                "descripcion": "C18, 150×4.6 mm, 5 µm",
                "fabricante": "Waters",
                "numero_parte": "WAT054275",
                "serial": "CL-2024-001",
                "numero_interno": "COL-00045",
            },
        ],
    }


def render_materiales():
    template = TPL / "materiales.docx"
    out = OUT / "materiales_renderizado.docx"
    ctx = context_materiales_demo()
    p = render_docxtpl(template, ctx, out)
    if p:
        print("✔ Materiales:", p)


# ============================================================
# SECCIÓN 4 — PRECISIÓN DEL SISTEMA (poda de columnas)
# ============================================================
TEMPLATE_ACTIVOS_MAX_PS = 5
AREA_DECIMALS = 3
RSD_DECIMALS = 2


def _pad_to_max(lst: List[Any], size: int, fill: Any) -> List[Any]:
    out = list(lst)
    if len(out) < size:
        out += [fill] * (size - len(out))
    return out[:size]


def build_context_precision_sistema(base: Dict[str, Any]) -> Dict[str, Any]:
    activos_in = base.get("activos", [])
    if not activos_in:
        raise ValueError("Faltan activos en base['activos'].")

    umbral_global = float(base.get("umbral_rsd", 2.0))
    umbrales_por_activo = base.get("umbrales_rsd_por_activo", None)
    if umbrales_por_activo is not None:
        umbrales_por_activo = list(map(float, umbrales_por_activo))
        if len(umbrales_por_activo) < len(activos_in):
            umbrales_por_activo += [umbral_global] * (
                len(activos_in) - len(umbrales_por_activo)
            )

    activos_proc, max_reps = [], 0
    for a in activos_in:
        areas = list(map(float, a.get("areas", [])))
        rsd = round(rsd_percent(areas), RSD_DECIMALS)
        activos_proc.append({"areas": areas, "rsd": rsd})
        max_reps = max(max_reps, len(areas))

    if max_reps == 0:
        raise ValueError("No hay réplicas para calcular precisión del sistema.")

    cumple_por_activo: List[bool] = []
    for idx, ap in enumerate(activos_proc):
        umbral_i = (
            umbrales_por_activo[idx]
            if umbrales_por_activo is not None
            else umbral_global
        )
        cumple_por_activo.append(ap["rsd"] <= round(float(umbral_i), RSD_DECIMALS))
    conclusion_global = "Cumple" if all(cumple_por_activo) else "No cumple"

    criterio_in = base.get("criterio", "").strip()
    criterio = (
        criterio_in
        if criterio_in
        else f"El %RSD de {max_reps} inyecciones replicadas de la solución estándar es ≤ {umbral_global:.1f}%."
    )

    activos_proc = _pad_to_max(
        activos_proc, TEMPLATE_ACTIVOS_MAX_PS, {"areas": [], "rsd": ""}
    )

    rows: List[Dict[str, Any]] = []
    for i in range(max_reps):
        row = {"replica": i + 1, "criterio": criterio, "conclusion": conclusion_global}
        for j in range(TEMPLATE_ACTIVOS_MAX_PS):
            val = (
                activos_proc[j]["areas"][i] if i < len(activos_proc[j]["areas"]) else ""
            )
            row[f"area_pico_{j+1}"] = "" if val == "" else f"{val:.{AREA_DECIMALS}f}"
        rows.append(row)

    rsd_vals = [ap["rsd"] for ap in activos_proc]
    rsd_vals = _pad_to_max(rsd_vals, TEMPLATE_ACTIVOS_MAX_PS, "")
    rsd_ctx = {
        f"RSD_precision_pico_{i+1}": (
            "" if rsd_vals[i] == "" else f"{rsd_vals[i]:.{RSD_DECIMALS}f}"
        )
        for i in range(TEMPLATE_ACTIVOS_MAX_PS)
    }
    return {
        "precision_sistema": rows,
        "conclusion_global_precision_sistema": conclusion_global,
        "criterio_precision_sistema": criterio,
        **rsd_ctx,
    }


def _find_precision_table(doc: Document):
    for t in doc.tables:
        if not t.rows:
            continue
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if len(hdr) >= 8 and hdr[0] == "Réplica" and hdr[-1] == "Conclusión":
            return t
    return None


def _delete_column(table, col_idx: int):
    for row in table.rows:
        cell = row.cells[col_idx]
        cell._tc.getparent().remove(cell._tc)


def prune_precision_sistema_columns(
    docx_path: Path, n_activos: int, max_activos: int = TEMPLATE_ACTIVOS_MAX_PS
):
    d = Document(str(docx_path))
    table = _find_precision_table(d)
    if table is None:
        d.save(str(docx_path))
        return
    cols_to_delete = list(range(1 + n_activos, 1 + max_activos))
    for col in sorted(cols_to_delete, reverse=True):
        _delete_column(table, col)
    d.save(str(docx_path))


def render_precision_sistema():
    template = TPL / "template_precision.docx"
    out = OUT / "precision_sistema.docx"
    base = {
        "umbral_rsd": 2.0,
        "activos": [
            {
                "nombre": "Activo1",
                "areas": [3115.0, 3120.1, 3118.2, 3116.9, 3110.2, 3119.7],
            },
            {"nombre": "Activo2", "areas": [944.0, 946.9, 943.9, 945.2, 945.1, 944.1]},
        ],
    }
    ctx = build_context_precision_sistema(base)
    p = render_docxtpl(template, ctx, out)
    if p:
        n_activos = min(len(base.get("activos", [])), TEMPLATE_ACTIVOS_MAX_PS)
        if n_activos < TEMPLATE_ACTIVOS_MAX_PS:
            prune_precision_sistema_columns(p, n_activos)
        print("✔ Precisión del sistema:", p)


# ============================================================
# SECCIÓN 5 — REPETIBILIDAD (método) — con poda de columnas
# ============================================================
TEMPLATE_ACTIVOS_MAX_REP = 5
_EPS = 1e-9  # tolerancia numérica


def _extract_threshold_from_text_rep(criterio_text: str) -> float | None:
    """Extrae un número (porcentaje) del texto del criterio, p.ej. '... ≤ 2.0%.' -> 2.0"""
    if not criterio_text:
        return None
    m = re.search(r"(\d+(?:[.,]\d+)?)\s*%?", criterio_text)
    if not m:
        return None
    num = m.group(1).replace(",", ".")
    try:
        return float(num)
    except ValueError:
        return None


def build_context_repetibilidad(base: Dict[str, Any]) -> Dict[str, Any]:
    """
    base = {
      "criterio": "El %RSD ... <= 2.0%.",
      "criterio_rsd_max": 2.0,    # opcional (si no, se intenta extraer del texto)
      "conclusion": "Cumple",     # opcional si no hay umbral
      "activos": [
        {"nombre": "A", "porcentajes": [floats...], "rsd": 0.23 (opcional)},
        ...
      ]
    }
    """
    criterio_text = base.get("criterio", "")
    criterio_rsd_max = base.get("criterio_rsd_max", None)
    if criterio_rsd_max is None:
        criterio_rsd_max = _extract_threshold_from_text_rep(criterio_text)

    activos_in = base.get("activos", [])
    if not activos_in:
        raise ValueError("Faltan activos en base['activos'].")

    activos_proc, max_reps = [], 0
    for a in activos_in:
        porcentajes = list(map(float, a.get("porcentajes", [])))
        rsd_val = a.get("rsd", None)
        if rsd_val is None:
            rsd_val = rsd_percent(porcentajes)
        activos_proc.append(
            {
                "nombre": a.get("nombre", ""),
                "porcentajes": [round(x, 2) for x in porcentajes],
                "rsd_val": float(rsd_val),
                "rsd_fmt": f"{float(rsd_val):.2f}",
            }
        )
        max_reps = max(max_reps, len(porcentajes))

    cumple_por_activo: List[str] = []
    all_cumple = True
    if criterio_rsd_max is not None:
        for ap in activos_proc:
            cumple = ap["rsd_val"] <= float(criterio_rsd_max) + _EPS
            cumple_por_activo.append("Cumple" if cumple else "No cumple")
            if not cumple:
                all_cumple = False
    else:
        cumple_por_activo = ["" for _ in activos_proc]
        all_cumple = None

    activos_proc = _pad_to_max(
        activos_proc,
        TEMPLATE_ACTIVOS_MAX_REP,
        {"nombre": "", "porcentajes": [], "rsd_val": 0.0, "rsd_fmt": ""},
    )
    cumple_por_activo = _pad_to_max(cumple_por_activo, TEMPLATE_ACTIVOS_MAX_REP, "")

    if all_cumple is True:
        conclusion_global = "Cumple"
    elif all_cumple is False:
        conclusion_global = "No cumple"
    else:
        conclusion_global = base.get("conclusion", "")

    rows: List[Dict[str, Any]] = []
    for i in range(max_reps):
        row = {
            "replica": i + 1,
            "criterio": criterio_text,
            "conclusion": conclusion_global,
        }
        for j in range(TEMPLATE_ACTIVOS_MAX_REP):
            val = (
                activos_proc[j]["porcentajes"][i]
                if i < len(activos_proc[j]["porcentajes"])
                else ""
            )
            row[f"porcentaje_pico_{j+1}"] = "" if val == "" else f"{val:.2f}"
        rows.append(row)

    rsd_ctx = {
        f"RSD_precision_pico_{i+1}": activos_proc[i]["rsd_fmt"]
        for i in range(TEMPLATE_ACTIVOS_MAX_REP)
    }
    conc_ctx = {
        f"CONC_precision_pico_{i+1}": cumple_por_activo[i]
        for i in range(TEMPLATE_ACTIVOS_MAX_REP)
    }
    extra_ctx = {
        "criterio_rsd_max_num": (
            "" if criterio_rsd_max is None else f"{float(criterio_rsd_max):.2f}"
        )
    }

    return {
        "precision_metodo": rows,
        **rsd_ctx,
        **conc_ctx,
        **extra_ctx,
        "conclusion_global": conclusion_global,
    }


def _find_repetibilidad_table(doc: Document):
    """Localiza la tabla por encabezado: empieza con 'Réplica' y termina con 'Conclusión'."""
    for t in doc.tables:
        if not t.rows:
            continue
        hdr = [c.text.strip() for c in t.rows[0].cells]
        # se esperan 8+ columnas: 0 Rep, 1..5 Activos, penúltima Criterio, última Conclusión
        if len(hdr) >= 8 and hdr[0] == "Réplica" and hdr[-1] == "Conclusión":
            return t
    return None


def prune_repetibilidad_columns(
    docx_path: str | Path, n_activos: int, max_activos: int = TEMPLATE_ACTIVOS_MAX_REP
):
    """Borra columnas Activo_(n_activos+1)..Activo_max en el DOCX ya renderizado."""
    p = Path(docx_path)
    doc = Document(str(p))
    table = _find_repetibilidad_table(doc)
    if table is None:
        doc.save(str(p))
        return
    cols_to_delete = list(range(1 + n_activos, 1 + max_activos))
    for col in sorted(cols_to_delete, reverse=True):  # derecha -> izquierda
        _delete_column(table, col)
    doc.save(str(p))


def render_repetibilidad():
    template = TPL / "template_repetibilidad.docx"
    out = OUT / "repetibilidad.docx"
    base = {
        "criterio": "El %RSD obtenido es menor o igual a 2.0%.",
        "criterio_rsd_max": 2.0,  # umbral explícito
        "activos": [
            {
                "nombre": "Activo1",
                "porcentajes": [100.10, 100.46, 100.34, 99.61, 100.17, 100.27],
            },
            {
                "nombre": "Activo2",
                "porcentajes": [199.60, 100.12, 100.43, 99.61, 100.17, 100.25],
            },
        ],
    }
    ctx = build_context_repetibilidad(base)
    p = render_docxtpl(template, ctx, out)
    if p:
        n_act_reales = min(len(base.get("activos", [])), TEMPLATE_ACTIVOS_MAX_REP)
        if n_act_reales < TEMPLATE_ACTIVOS_MAX_REP:
            prune_repetibilidad_columns(p, n_act_reales, TEMPLATE_ACTIVOS_MAX_REP)
        print("✔ Repetibilidad:", p)


# ============================================================
# SECCIÓN 6 — PRECISIÓN INTERMEDIA (poda con gridSpan)
# ============================================================
TEMPLATE_ACTIVOS_MAX_PI = 5  # pares AN1/AN2


def _rsd(vals: List[float]) -> float:
    vals = [float(v) for v in vals if v is not None]
    if len(vals) < 2:
        return 0.0
    m = mean(vals)
    if m == 0:
        return 0.0
    return 100.0 * stdev(vals) / m


def _pad5_precint(acts: List[Optional[Dict[str, Any]]]) -> List[Dict[str, str]]:
    out = []
    for i in range(TEMPLATE_ACTIVOS_MAX_PI):
        a = acts[i] if i < len(acts) and acts[i] else {}
        out.append({"an1": fmt2(a.get("an1")), "an2": fmt2(a.get("an2"))})
    return out


def _n_activos_con_datos_precint(
    filas_input: List[Dict[str, Any]], max_activos: int = TEMPLATE_ACTIVOS_MAX_PI
) -> int:
    n = 0
    for i in range(max_activos):
        hay = False
        for f in filas_input:
            acts = f.get("activos", [])
            if i < len(acts) and acts[i]:
                a = acts[i]
                if a.get("an1") is not None or a.get("an2") is not None:
                    hay = True
                    break
        if hay:
            n = i + 1
        else:
            break
    return n


def construir_contexto_precision_intermedia(
    filas_input: List[Dict[str, Any]], criterio_max: float = 2.0
) -> Dict[str, Any]:
    if not filas_input:
        raise ValueError("filas_input vacío")
    filas_ord = sorted(filas_input, key=lambda d: int(d["replica"]))
    filas = [
        {"replica": int(f["replica"]), "activos": _pad5_precint(f.get("activos", []))}
        for f in filas_ord
    ]

    rsd_por_activo = []
    for i in range(TEMPLATE_ACTIVOS_MAX_PI):
        pool = []
        for f in filas_ord:
            acts = f.get("activos", [])
            if i < len(acts) and acts[i]:
                a = acts[i]
                if a.get("an1") is not None:
                    pool.append(float(a["an1"]))
                if a.get("an2") is not None:
                    pool.append(float(a["an2"]))
        rsd_por_activo.append(round(_rsd(pool), 2) if len(pool) >= 2 else "")

    definidos = [v for v in rsd_por_activo if isinstance(v, (int, float))]
    conclusion_global = (
        "Cumple"
        if (definidos and all(v <= criterio_max for v in definidos))
        else "No cumple"
    )
    criterio_txt = f"El %RSD obtenido es menor o igual a {criterio_max:.1f}%."
    return {
        "filas": filas,
        "rsd_por_activo": rsd_por_activo,
        "conclusion_global": conclusion_global,
        "criterio_txt": criterio_txt,
    }


# --- Helpers gridSpan (versión PI, nombres únicos) ---
def _get_span_pi(tc) -> int:
    tcPr = tc.tcPr
    if tcPr is not None:
        gridSpan = tcPr.find(qn("w:gridSpan"))
        if gridSpan is not None and gridSpan.get(qn("w:val")):
            try:
                return int(gridSpan.get(qn("w:val")))
            except Exception:
                return 1
    return 1


def _set_span_pi(tc, span: int):
    tcPr = tc.get_or_add_tcPr()
    gridSpan = tcPr.find(qn("w:gridSpan"))
    if span <= 1:
        if gridSpan is not None:
            tcPr.remove(gridSpan)
    else:
        if gridSpan is None:
            gridSpan = OxmlElement("w:gridSpan")
            tcPr.append(gridSpan)
        gridSpan.set(qn("w:val"), str(span))


def _row_grid_map_pi(tr):
    out = []
    cur = 0
    for tc in tr.tc_lst:
        sp = _get_span_pi(tc)
        out.append((tc, cur, sp))
        cur += sp
    return out


def _ncols_from_first_row_pi(table) -> int:
    tr0 = table._tbl.tr_lst[0]
    return sum(_get_span_pi(tc) for tc in tr0.tc_lst)


def _delete_grid_col_pi(table, col_idx: int):
    for tr in table._tbl.tr_lst:
        target = None
        for tc, start, span in _row_grid_map_pi(tr):
            if start <= col_idx < start + span:
                target = (tc, span)
                break
        if target is None:
            continue
        tc, span = target
        if span > 1:
            _set_span_pi(tc, span - 1)
        else:
            tr.remove(tc)


def prune_activos_trailing_empty(docx_path: Path, n_activos_keep: int) -> Path:
    p = Path(docx_path)
    d = Document(str(p))
    table = None
    for t in d.tables:
        if not t.rows:
            continue
        hdr = [c.text.strip() for c in t.rows[0].cells]
        if (
            len(hdr) >= 3
            and hdr[0] == "Réplica"
            and ("Conclusión" in hdr[-1] or hdr[-1] == "Conclusión")
        ):
            table = t
            break
    if table is None:
        d.save(str(p))
        return p

    n_total_cols = _ncols_from_first_row_pi(table)  # = 1 + 2*max_act + 2
    max_act = (n_total_cols - 3) // 2
    n_activos_keep = max(0, min(n_activos_keep, max_act))

    cols_to_delete = []
    for i in range(n_activos_keep, max_act):
        an1 = 1 + 2 * i
        an2 = 1 + 2 * i + 1
        cols_to_delete.extend([an1, an2])
    for col in sorted(cols_to_delete, reverse=True):
        _delete_grid_col_pi(table, col)

    d.save(str(p))
    return p


def render_precision_intermedia():
    template = TPL / "template_precision_intermedia_tr.docx"
    out = OUT / "precision_intermedia_renderizada.docx"
    filas_input = [
        {
            "replica": 1,
            "activos": [{"an1": 99.10, "an2": 100.05}, {"an1": 98.70, "an2": 99.90}],
        },
        {
            "replica": 2,
            "activos": [{"an1": 100.20, "an2": 99.80}, {"an1": 100.10, "an2": 99.95}],
        },
        {
            "replica": 3,
            "activos": [{"an1": 99.80, "an2": 100.10}, {"an1": 100.30, "an2": 100.05}],
        },
    ]
    ctx = construir_contexto_precision_intermedia(filas_input, criterio_max=2.0)
    p = render_docxtpl(template, ctx, out)
    if p:
        n_keep = _n_activos_con_datos_precint(
            filas_input, max_activos=TEMPLATE_ACTIVOS_MAX_PI
        )
        prune_activos_trailing_empty(p, n_keep)
        print("✔ Precisión intermedia:", p)


# ============================================================
# SECCIÓN 7 — ESTABILIDAD DE SOLUCIONES
# ============================================================
def evaluar_cumple(
    area_t: Optional[float],
    area_t0: Optional[float],
    delta_max_pct: float = 2.0,
    similitud_lo: float = 98.0,
    similitud_hi: float = 102.0,
) -> Dict[str, Any]:
    if area_t is None or area_t0 is None or area_t0 == 0:
        return {"di_pct": "", "delta_pct": "", "cumple": None}
    di = 100.0 * (area_t / area_t0)
    delta = 100.0 * (area_t - area_t0) / area_t0
    cumple = (abs(delta) <= delta_max_pct) or (similitud_lo <= di <= similitud_hi)
    return {"di_pct": round(di, 2), "delta_pct": round(delta, 2), "cumple": cumple}


def construir_contexto_estabilidad(
    activos_input: List[Dict[str, Any]],
    tiempos: List[str] = ["T0", "T1", "T2", "T3"],
    condiciones_por_tiempo: Dict[str, List[str]] = {
        "T1": ["C1", "C2"],
        "T2": ["C1", "C2"],
        "T3": ["C1", "C2"],
    },
    criterio_delta_max_pct: float = 2.0,
    criterio_sim_lo: float = 98.0,
    criterio_sim_hi: float = 102.0,
) -> Dict[str, Any]:
    filas = []
    glob_keys = ["T1C1", "T1C2", "T2C1", "T2C2", "T3C1", "T3C2"]
    concl_flags = {k: [] for k in glob_keys}

    for a in activos_input:
        nombre = str(a["nombre"])
        m = a.get("medidas", {})
        t0_raw = m.get("T0")
        if isinstance(t0_raw, list):
            t0_val = safe_mean(t0_raw)
        elif t0_raw is None:
            t0_val = None
        else:
            t0_val = float(t0_raw)

        a_row: Dict[str, Any] = {"nombre": nombre}
        a_row["area_T0"] = fmt2(t0_val)

        for t in ["T1", "T2", "T3"]:
            conds = condiciones_por_tiempo.get(t, [])
            for c in conds:
                key = f"{t}{c}"
                raw = m.get(t, {}).get(c)
                if isinstance(raw, list):
                    area = safe_mean(raw)
                else:
                    area = float(raw) if raw is not None else None
                a_row[f"area_{t}{c}"] = fmt2(area)
                res = evaluar_cumple(
                    area,
                    t0_val,
                    criterio_delta_max_pct,
                    criterio_sim_lo,
                    criterio_sim_hi,
                )
                a_row[f"di_{t}{c}"] = (
                    "" if res["di_pct"] == "" else f"{res['di_pct']:.2f}"
                )
                if res["cumple"] is not None:
                    concl_flags[key].append(bool(res["cumple"]))

        for key in glob_keys:
            a_row.setdefault(f"area_{key}", "")
            a_row.setdefault(f"di_{key}", "")
        filas.append(a_row)

    concl_global: Dict[str, str] = {}
    for key, flags in concl_flags.items():
        if not flags:
            concl_global[key] = "No aplica"
        else:
            concl_global[key] = "Cumple" if all(flags) else "No cumple"

    criterio_txt = (
        f"La solución estándar y la solución muestra se consideran estables si "
        f"el porcentaje de cambio es ≤ {criterio_delta_max_pct:.1f}% "
        f"o si el porcentaje de similitud se mantiene entre {criterio_sim_lo:.1f}% y {criterio_sim_hi:.1f}% "
        f"respecto a T0."
    )
    return {
        "activos": filas,
        "criterio_txt": criterio_txt,
        "concl_T1C1": concl_global["T1C1"],
        "concl_T1C2": concl_global["T1C2"],
        "concl_T2C1": concl_global["T2C1"],
        "concl_T2C2": concl_global["T2C2"],
        "concl_T3C1": concl_global["T3C1"],
        "concl_T3C2": concl_global["T3C2"],
    }


def render_estabilidad_soluciones():
    template = TPL / "estabilidad_tabla_tr.docx"
    out = OUT / "estabilidad_renderizada.docx"
    activos_input = [
        {
            "nombre": "Activo_1",
            "medidas": {
                "T0": [1000.0, 1001.5, 998.5],
                "T1": {"C1": [995.0, 996.0], "C2": [1004.0, 1003.5]},
                "T2": {"C1": [985.0, 986.0], "C2": [1001.0, 1002.0]},
                "T3": {"C1": [980.0, 981.0], "C2": [997.0, 998.0]},
            },
        },
        {
            "nombre": "Activo_2",
            "medidas": {
                "T0": [2000.0, 1999.0, 2001.0],
                "T1": {"C1": [1990.0, 1992.0], "C2": [2010.0, 2008.0]},
                "T2": {"C1": [1988.0, 1987.0], "C2": [2011.0, 2012.0]},
                "T3": {"C1": [1979.0, 1981.0], "C2": [2005.0, 2006.0]},
            },
        },
    ]
    ctx = construir_contexto_estabilidad(activos_input)
    p = render_docxtpl(template, ctx, out)
    if p:
        print("✔ Estabilidad de soluciones:", p)


# ============================================================
# SECCIÓN 8 — ESTABILIDAD DE FASE MÓVIL (poda por bloque de tiempos)
# ============================================================
MAX_ACT_FM = 5


def _safe_at(seq, i):
    if isinstance(seq, list) and 0 <= i < len(seq):
        return seq[i]
    return None


def _parse_float(s):
    try:
        return float(s)
    except Exception:
        return None


def construir_contexto_fase_movil(
    nombres_activos: List[str],
    replicas_input: List[Dict[str, Any]],
    rsd_t0: Optional[List[Optional[float]]] = None,
    rsd_t1: Optional[List[Optional[float]]] = None,
    rsd_t2: Optional[List[Optional[float]]] = None,
    asim_t0: Optional[List[Optional[float]]] = None,
    asim_t1: Optional[List[Optional[float]]] = None,
    asim_t2: Optional[List[Optional[float]]] = None,
    exact_t0: Optional[List[Optional[float]]] = None,
    exact_t1: Optional[List[Optional[float]]] = None,
    exact_t2: Optional[List[Optional[float]]] = None,
    resol_T0: Optional[float] = None,
    resol_T1: Optional[float] = None,
    resol_T2: Optional[float] = None,
    concl_T1: str = "Cumple",
    concl_T2: str = "Cumple",
    criterio_txt: str = (
        "La fase móvil se considera estable si el rango de variación del tiempo de retención es ±3 minutos "
        "respecto al análisis inicial, no hay picos fantasmas/división, "
        "cada tiempo cumple adecuabilidad del sistema (RSD ≤ 2.0%, asimetría ≤ 1.6, resolución ≥ 5.0) "
        "y el factor de exactitud se mantiene en 98.0–102.0%."
    ),
) -> Dict[str, Any]:
    n_act = min(len(nombres_activos), MAX_ACT_FM)
    headers = [nombres_activos[i] if i < n_act else "" for i in range(MAX_ACT_FM)]
    reps_sorted = sorted(replicas_input, key=lambda x: int(x.get("num", 0)))

    filas = []
    for rep in reps_sorted:
        row = {"num": int(rep.get("num", 0))}
        for key in ["t0", "t1", "t2"]:
            raw = rep.get(key, [])
            vals_fmt = [fmt2(_safe_at(raw, i)) for i in range(MAX_ACT_FM)]
            row[key] = vals_fmt
        filas.append(row)

    tr_mean_t0, tr_mean_t1, tr_mean_t2 = [], [], []
    for i in range(MAX_ACT_FM):
        mt0 = safe_mean([_safe_at(rep.get("t0", []), i) for rep in replicas_input])
        mt1 = safe_mean([_safe_at(rep.get("t1", []), i) for rep in replicas_input])
        mt2 = safe_mean([_safe_at(rep.get("t2", []), i) for rep in replicas_input])
        tr_mean_t0.append(fmt2(mt0))
        tr_mean_t1.append(fmt2(mt1))
        tr_mean_t2.append(fmt2(mt2))

    def delta_block(tr_t_list, tr_0_list):
        diffs = []
        for a_str, b_str in zip(tr_t_list, tr_0_list):
            at = _parse_float(a_str)
            a0 = _parse_float(b_str)
            if at is not None and a0 is not None:
                diffs.append(abs(at - a0))
        return "" if not diffs else f"{mean(diffs):.3f}"

    deltaT_T1 = delta_block(tr_mean_t1, tr_mean_t0)
    deltaT_T2 = delta_block(tr_mean_t2, tr_mean_t0)

    def padlist(x):
        if x is None:
            return [""] * MAX_ACT_FM
        out = []
        for i in range(MAX_ACT_FM):
            v = x[i] if i < len(x) else None
            out.append("" if v is None else fmt2(v))
        return out

    return {
        "headers_activos": headers,
        "replicas": filas,
        "rsd_t0": padlist(rsd_t0),
        "rsd_t1": padlist(rsd_t1),
        "rsd_t2": padlist(rsd_t2),
        "asim_t0": padlist(asim_t0),
        "asim_t1": padlist(asim_t1),
        "asim_t2": padlist(asim_t2),
        "exact_t0": padlist(exact_t0),
        "exact_t1": padlist(exact_t1),
        "exact_t2": padlist(exact_t2),
        "tr_mean_t0": tr_mean_t0,
        "tr_mean_t1": tr_mean_t1,
        "tr_mean_t2": tr_mean_t2,
        "resol_T0": "" if resol_T0 is None else fmt2(resol_T0),
        "resol_T1": "" if resol_T1 is None else fmt2(resol_T1),
        "resol_T2": "" if resol_T2 is None else fmt2(resol_T2),
        "deltaT_T1": deltaT_T1,
        "deltaT_T2": deltaT_T2,
        "criterio_txt": criterio_txt,
        "concl_T1": concl_T1,
        "concl_T2": concl_T2,
    }


# Helpers gridSpan (versión FM, nombres únicos)
def _get_span_fm(tc) -> int:
    tcPr = tc.tcPr
    if tcPr is not None:
        g = tcPr.find(qn("w:gridSpan"))
        if g is not None and g.get(qn("w:val")):
            try:
                return int(g.get(qn("w:val")))
            except Exception:
                return 1
    return 1


def _set_span_fm(tc, span: int):
    tcPr = tc.get_or_add_tcPr()
    g = tcPr.find(qn("w:gridSpan"))
    if span <= 1:
        if g is not None:
            tcPr.remove(g)
    else:
        if g is None:
            g = OxmlElement("w:gridSpan")
            tcPr.append(g)
        g.set(qn("w:val"), str(span))


def _row_grid_map_fm(tr):
    out = []
    cur = 0
    for tc in tr.tc_lst:
        sp = _get_span_fm(tc)
        out.append((tc, cur, sp))
        cur += sp
    return out


def _ncols_from_first_row_fm(table) -> int:
    tr0 = table._tbl.tr_lst[0]
    return sum(_get_span_fm(tc) for tc in tr0.tc_lst)


def _delete_grid_col_fm(table, col_idx: int):
    for tr in table._tbl.tr_lst:
        target = None
        for tc, start, span in _row_grid_map_fm(tr):
            if start <= col_idx < start + span:
                target = (tc, span)
                break
        if target is None:
            continue
        tc, span = target
        if span > 1:
            _set_span_fm(tc, span - 1)
        else:
            tr.remove(tc)


def prune_activos_trailing_fm(
    docx_path: Path, n_act_keep: int, n_times: int = 3
) -> Path:
    p = Path(docx_path)
    d = Document(str(p))
    if not d.tables:
        d.save(str(p))
        return p
    table = d.tables[0]
    total_cols = _ncols_from_first_row_fm(table)
    max_act = (total_cols - 1) // n_times
    cols_to_delete = []
    for t in range(n_times):
        start = 1 + t * max_act
        for i in range(n_act_keep, max_act):
            cols_to_delete.append(start + i)
    for col in sorted(cols_to_delete, reverse=True):
        _delete_grid_col_fm(table, col)
    d.save(str(p))
    return p


def render_fase_movil():
    template = TPL / "fase_movil_tabla_tr.docx"
    out = OUT / "fase_movil_renderizada.docx"
    nombres_activos = ["Activo_1", "Activo_2"]
    replicas_input = [
        {"num": 1, "t0": [3.094, 6.191], "t1": [3.084, 6.173], "t2": [3.093, 6.510]},
        {"num": 2, "t0": [3.096, 6.188], "t1": [3.085, 6.170], "t2": [3.091, 6.511]},
        {"num": 3, "t0": [3.095, 6.192], "t1": [3.083, 6.174], "t2": [3.092, 6.509]},
        {"num": 4, "t0": [3.097, 6.189], "t1": [3.086, 6.171], "t2": [3.094, 6.512]},
        {"num": 5, "t0": [3.094, 6.190], "t1": [3.084, 6.172], "t2": [3.093, 6.510]},
    ]
    ctx = construir_contexto_fase_movil(
        nombres_activos,
        replicas_input,
        rsd_t0=[0.0, 0.0],
        rsd_t1=[0.0, 0.1],
        rsd_t2=[0.0, 0.2],
        asim_t0=[1.10, 1.10],
        asim_t1=[1.10, 1.10],
        asim_t2=[1.10, 1.10],
        exact_t0=[100.1, 100.0],
        exact_t1=[100.0, 100.0],
        exact_t2=[99.9, 99.8],
        resol_T0=18.1,
        resol_T1=18.1,
        resol_T2=18.8,
        concl_T1="Cumple",
        concl_T2="Cumple",
    )
    p = render_docxtpl(template, ctx, out)
    if p:
        prune_activos_trailing_fm(p, n_act_keep=len(nombres_activos), n_times=3)
        print("✔ Fase móvil:", p)


# ============================================================
# SECCIÓN 9 — ROBUSTEZ (contenido promedio y |di|)
# ============================================================
def render_robustez_contenido_di():
    template = TPL / "robustez_contenido_di_template.docx"
    out = OUT / "robustez_contenido_di_renderizado.docx"
    context = {
        "compuestos": [
            {
                "nombre": "Activo_1",
                "param_rows": [
                    {
                        "parametro": "Contenido promedio (%)",
                        "valores": [
                            fmt2(v)
                            for v in [
                                100.90,
                                101.14,
                                101.83,
                                100.98,
                                100.94,
                                100.95,
                                100.76,
                                100.79,
                                100.73,
                            ]
                        ],
                    },
                    {
                        "parametro": "|di| (%)",
                        "valores": [
                            "No aplica",
                            fmt2(0.2),
                            fmt2(0.9),
                            fmt2(0.1),
                            fmt2(0.0),
                            fmt2(0.0),
                            fmt2(0.1),
                            fmt2(0.1),
                            fmt2(0.2),
                        ],
                    },
                ],
            },
            {
                "nombre": "Activo_2",
                "param_rows": [
                    {
                        "parametro": "Contenido promedio (%)",
                        "valores": [
                            fmt2(v)
                            for v in [
                                99.31,
                                100.51,
                                100.39,
                                99.31,
                                99.34,
                                99.51,
                                98.77,
                                97.69,
                                99.55,
                            ]
                        ],
                    },
                    {
                        "parametro": "|di| (%)",
                        "valores": [
                            "No aplica",
                            fmt2(1.2),
                            fmt2(1.1),
                            fmt2(0.0),
                            fmt2(0.0),
                            fmt2(0.2),
                            fmt2(0.5),
                            fmt2(1.6),
                            fmt2(1.2),
                        ],
                    },
                ],
            },
        ],
        "criterio_txt": (
            "Las especificaciones del Test de Adecuabilidad del Sistema (SST) deben cumplirse en todas las condiciones analizadas. "
            "El porcentaje de cambio |di| obtenido en cada condición evaluada debe ser ≤ 2.0% respecto a la condición nominal."
        ),
        "conclusiones": ["No aplica"] + ["Cumple"] * 8,
    }
    p = render_docxtpl(template, context, out)
    if p:
        print("✔ Robustez (contenido y |di|):", p)


# ============================================================
# MAIN — Ejecuta todo en serie, sin caerse si falta algo
# ============================================================
if __name__ == "__main__":
    print("\n=== Renderizador unificado de plantillas de validación ===\n")

    try:
        render_linealidad()
    except Exception as e:
        print("✖ Linealidad: ", e)

    try:
        render_exactitud_multi()
    except Exception as e:
        print("✖ Exactitud multi: ", e)

    try:
        render_materiales()
    except Exception as e:
        print("✖ Materiales: ", e)

    try:
        render_precision_sistema()
    except Exception as e:
        print("✖ Precisión del sistema: ", e)

    try:
        render_repetibilidad()
    except Exception as e:
        print("✖ Repetibilidad: ", e)

    try:
        render_precision_intermedia()
    except Exception as e:
        print("✖ Precisión intermedia: ", e)

    try:
        render_estabilidad_soluciones()
    except Exception as e:
        print("✖ Estabilidad de soluciones: ", e)

    try:
        render_fase_movil()
    except Exception as e:
        print("✖ Fase móvil: ", e)

    try:
        render_robustez_contenido_di()
    except Exception as e:
        print("✖ Robustez (contenido |di|): ", e)

    print("\n>> Listo. Revisa la carpeta de salida:\n   ", OUT)
